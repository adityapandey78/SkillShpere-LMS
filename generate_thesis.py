"""
SkillSphere LMS — B.Tech Thesis Generator
Generates a 26–30 page DOCX following NIAMT thesis template.
"""

import os, io
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrow
import matplotlib.patheffects as pe
import numpy as np

OUT = r"d:\Coding\Development\ProjctsOP\SkillShpere-LMS\SkillSphere_BThesis.docx"

# ─────────────────────────────────────────────────────────────
# STYLE HELPERS
# ─────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_spacing(para, before=6, after=6, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)

def heading(doc, text, level=1, caps=True, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper() if caps else text)
    set_font(run, size=size, bold=True)
    para_spacing(p, before=12, after=6)
    return p

def section_heading(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    para_spacing(p, before=10, after=4)
    return p

def body(doc, text, indent=0, justify=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=12)
    para_spacing(p, before=0, after=6)
    return p

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=10, italic=True)
    para_spacing(p, before=2, after=10)

def spacer(doc, pts=8):
    """Blank paragraph used to prevent image/table crowding."""
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=pts)
    return p

def insert_figure(doc, buf, width_in, caption_text):
    """Insert an image with guaranteed clear space above and below."""
    spacer(doc, pts=10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_in))
    para_spacing(p, before=4, after=4)
    caption(doc, caption_text)
    spacer(doc, pts=6)

def set_margins(doc, top=1, bottom=1, left=1.25, right=1):
    for sec in doc.sections:
        sec.top_margin    = Inches(top)
        sec.bottom_margin = Inches(bottom)
        sec.left_margin   = Inches(left)
        sec.right_margin  = Inches(right)

def add_table_style(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

# ─────────────────────────────────────────────────────────────
# DIAGRAM GENERATORS  (return BytesIO PNG)
# ─────────────────────────────────────────────────────────────

BLUE   = "#2C5F8A"
LBLUE  = "#4A90C4"
GREEN  = "#27AE60"
ORANGE = "#E67E22"
PURPLE = "#8E44AD"
GRAY   = "#7F8C8D"
WHITE  = "#FFFFFF"
DKBLUE = "#1A3A5C"

def _save(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight",
                facecolor="white")
    buf.seek(0)
    plt.close(fig)
    return buf


def diag_three_tier():
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10); ax.set_ylim(0, 7); ax.axis("off")
    fig.patch.set_facecolor("white")

    tiers = [
        (0.3, 5.5, 9.4, 1.3, BLUE,   "TIER 1 — PRESENTATION LAYER",
         ["React 18 + Vite", "Radix UI + Tailwind CSS",
          "Context API (Auth / Student / Instructor)",
          "services/index.js  ·  Framer Motion"]),
        (0.3, 3.5, 9.4, 1.3, PURPLE, "TIER 2 — APPLICATION LAYER",
         ["Express.js REST API  (Node.js 18)",
          "JWT Middleware  ·  Multer Upload  ·  CORS",
          "Controllers: auth | student | instructor | ai",
          "Helpers: Gemini · Cloudinary · PayPal"]),
        (0.3, 1.5, 9.4, 1.3, GREEN,  "TIER 3 — DATA LAYER",
         ["MongoDB Atlas (7 Collections)",
          "Users · Courses · Orders · Quizzes · Attempts",
          "CourseProgress · StudentCourses"]),
    ]

    for x, y, w, h, color, title, items in tiers:
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle="round,pad=0.05",
                              linewidth=1.5,
                              edgecolor=color, facecolor=color + "22")
        ax.add_patch(rect)
        ax.text(x + w/2, y + h - 0.18, title,
                ha="center", va="top", fontsize=9.5, fontweight="bold",
                color=color)
        row = "   ·   ".join(items[:2]) + "\n" + "   ·   ".join(items[2:])
        ax.text(x + w/2, y + h/2 - 0.15, row,
                ha="center", va="center", fontsize=8,
                color="#333333", linespacing=1.6)

    # Arrows between tiers
    for ya, yb in [(5.5, 4.8), (3.5, 2.8)]:
        ax.annotate("", xy=(5, yb), xytext=(5, ya),
                    arrowprops=dict(arrowstyle="<->",
                                   color=GRAY, lw=2))
        ax.text(5.2, (ya + yb)/2, "HTTPS + Mongoose ODM",
                fontsize=7.5, color=GRAY, va="center")

    # External services row
    ext = [(1.0, 0.2, LBLUE,   "MongoDB\nAtlas"),
           (3.2, 0.2, ORANGE,  "Cloudinary\nCDN"),
           (5.4, 0.2, GREEN,   "Google\nGemini 2.0"),
           (7.6, 0.2, "#C0392B","PayPal\nSandbox")]
    for ex, ey, ec, et in ext:
        r = FancyBboxPatch((ex, ey), 1.8, 0.9,
                           boxstyle="round,pad=0.05",
                           edgecolor=ec, facecolor=ec + "33", lw=1.5)
        ax.add_patch(r)
        ax.text(ex + 0.9, ey + 0.45, et, ha="center", va="center",
                fontsize=8, fontweight="bold", color=ec)
        ax.annotate("", xy=(ex + 0.9, 1.5),
                    xytext=(ex + 0.9, ey + 0.9),
                    arrowprops=dict(arrowstyle="->",
                                   color=GRAY, lw=1.2, linestyle="dashed"))

    ax.set_title("Figure 3.1 — SkillSphere Three-Tier Architecture",
                 fontsize=11, fontweight="bold", pad=10)
    return _save(fig)


def diag_er():
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_xlim(0, 12); ax.set_ylim(0, 8); ax.axis("off")
    fig.patch.set_facecolor("white")

    entities = {
        "User":           (0.4, 5.5, 2.2, 1.8, BLUE),
        "Course":         (4.5, 5.5, 2.2, 1.8, BLUE),
        "Order":          (8.5, 5.5, 2.2, 1.8, ORANGE),
        "StudentCourses": (0.4, 2.2, 2.2, 1.8, GREEN),
        "CourseProgress": (4.5, 2.2, 2.2, 1.8, GREEN),
        "Quiz":           (8.5, 2.2, 2.2, 1.8, PURPLE),
        "QuizAttempt":    (8.5, 0.1, 2.2, 1.8, PURPLE),
    }

    fields = {
        "User":           ["_id (PK)", "userName", "userEmail", "password", "role"],
        "Course":         ["_id (PK)", "instructorId", "title", "pricing", "curriculum[]"],
        "Order":          ["_id (PK)", "userId", "courseId", "paymentStatus", "orderStatus"],
        "StudentCourses": ["_id (PK)", "userId", "courses[]"],
        "CourseProgress": ["_id (PK)", "userId", "courseId", "completed", "lecturesProgress[]"],
        "Quiz":           ["_id (PK)", "courseId (unique)", "config{}", "groups[]"],
        "QuizAttempt":    ["_id (PK)", "userId", "courseId", "score", "passed"],
    }

    centers = {}
    for name, (x, y, w, h, color) in entities.items():
        # Header
        header = FancyBboxPatch((x, y + h - 0.35), w, 0.35,
                                boxstyle="round,pad=0.0",
                                edgecolor=color, facecolor=color, lw=0)
        ax.add_patch(header)
        body_r = FancyBboxPatch((x, y), w, h - 0.35,
                                boxstyle="round,pad=0.0",
                                edgecolor=color, facecolor=color + "18", lw=1.5)
        ax.add_patch(body_r)
        ax.text(x + w/2, y + h - 0.17, name,
                ha="center", va="center",
                fontsize=8.5, fontweight="bold", color=WHITE)
        for i, f in enumerate(fields[name]):
            ax.text(x + 0.1, y + h - 0.5 - i * 0.26,
                    f, fontsize=7.2, color="#222222", va="top")
        centers[name] = (x + w/2, y + h/2)

    # Relationships
    rels = [
        ("User", "Course",         "1  teaches  *"),
        ("User", "StudentCourses", "1  owns  1"),
        ("Course", "CourseProgress","1  tracked by  *"),
        ("Course", "Quiz",         "1  has  1"),
        ("Quiz", "QuizAttempt",    "1  attempted  *"),
        ("User", "Order",          "1  places  *"),
    ]
    for a, b, label in rels:
        x1, y1 = centers[a]
        x2, y2 = centers[b]
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>",
                                   color=GRAY, lw=1.3,
                                   connectionstyle="arc3,rad=0.0"))
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx, my + 0.15, label,
                fontsize=6.5, color=GRAY,
                ha="center", bbox=dict(facecolor="white", alpha=0.7, lw=0))

    ax.set_title("Figure 3.2 — Entity Relationship Diagram (SkillSphere Database)",
                 fontsize=11, fontweight="bold", pad=10)
    return _save(fig)


def diag_ai_pipeline():
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6); ax.axis("off")
    fig.patch.set_facecolor("white")

    steps = [
        (0.3, 3.8, 1.8, 1.0, BLUE,   "Instructor\nInput",    "Topic, level,\naudiece"),
        (2.6, 3.8, 1.8, 1.0, LBLUE,  "AI Prompt\nBuilder",   "Structured\nprompt"),
        (4.9, 3.8, 1.8, 1.0, PURPLE, "Gemini 2.0\nFlash",    "JSON / SSE\nresponse"),
        (7.2, 3.8, 1.8, 1.0, GREEN,  "Response\nParser",     "Schema\nvalidation"),
        (9.1, 3.8, 1.8, 1.0, ORANGE, "MongoDB\nAtlas",       "Persist &\nreturn"),
    ]

    for (x, y, w, h, col, title, sub) in steps:
        r = FancyBboxPatch((x, y), w, h,
                           boxstyle="round,pad=0.08",
                           edgecolor=col, facecolor=col + "33", lw=2)
        ax.add_patch(r)
        ax.text(x+w/2, y+h-0.22, title,
                ha="center", va="top", fontsize=8, fontweight="bold", color=col)
        ax.text(x+w/2, y+0.25, sub,
                ha="center", va="bottom", fontsize=7, color="#444")

    # Arrows
    for x in [2.1, 4.4, 6.7, 8.9]:
        ax.annotate("", xy=(x + 0.5, 4.3), xytext=(x, 4.3),
                    arrowprops=dict(arrowstyle="-|>", color=GRAY, lw=1.8))

    # Three AI feature rows
    features = [
        (1.0, 2.4, 2.5, BLUE,   "Course Outline\nGeneration",
         "JSON mode · 3500 tokens\nFlattened lecture curriculum"),
        (4.25, 2.4, 2.5, PURPLE, "Quiz Generation",
         "MCQ per lecture group\nDifficulty distribution config"),
        (7.5, 2.4, 2.5, GREEN,  "AI Tutor Chat\n(Streaming SSE)",
         "Course-scoped context\nReal-time token streaming"),
    ]
    for (x, y, w, col, title, desc) in features:
        r = FancyBboxPatch((x, y), w, 1.1,
                           boxstyle="round,pad=0.06",
                           edgecolor=col, facecolor=col + "22", lw=1.5)
        ax.add_patch(r)
        ax.text(x+w/2, y+0.85, title,
                ha="center", fontsize=8, fontweight="bold", color=col)
        ax.text(x+w/2, y+0.22, desc,
                ha="center", fontsize=7, color="#444")
        cx = x + w/2
        ax.annotate("", xy=(cx, 3.8), xytext=(cx, 3.5),
                    arrowprops=dict(arrowstyle="->", color=col, lw=1.3))

    ax.text(5.5, 0.7, "All three features share the same Gemini helper — differing only in "
            "prompt template, token limit, and response mode (JSON vs. SSE stream).",
            ha="center", fontsize=8, style="italic", color=GRAY,
            wrap=True)

    ax.set_title("Figure 3.3 — AI Feature Pipeline (Gemini 2.0 Flash Integration)",
                 fontsize=11, fontweight="bold", pad=10)
    return _save(fig)


def diag_usecase():
    fig, ax = plt.subplots(figsize=(11, 8))
    ax.set_xlim(0, 11); ax.set_ylim(0, 8); ax.axis("off")
    fig.patch.set_facecolor("white")

    # System boundary
    rect = FancyBboxPatch((1.5, 0.3), 8, 7.2,
                          boxstyle="round,pad=0.1",
                          edgecolor="#AAAAAA", facecolor="#FAFAFA", lw=2, zorder=0)
    ax.add_patch(rect)
    ax.text(5.5, 7.35, "SkillSphere LMS System",
            ha="center", fontsize=10, fontweight="bold", color="#333")

    def actor(x, y, label, color=BLUE):
        # Stick figure
        ax.plot(x, y+0.55, 'o', ms=10, color=color, zorder=5)
        ax.plot([x, x], [y+0.45, y+0.05], color=color, lw=2, zorder=5)
        ax.plot([x-0.2, x+0.2], [y+0.3, y+0.3], color=color, lw=2, zorder=5)
        ax.plot([x, x-0.18], [y+0.05, y-0.25], color=color, lw=2, zorder=5)
        ax.plot([x, x+0.18], [y+0.05, y-0.25], color=color, lw=2, zorder=5)
        ax.text(x, y-0.45, label, ha="center", fontsize=8.5,
                fontweight="bold", color=color)

    def usecase(x, y, text, color=BLUE):
        ell = mpatches.Ellipse((x, y), 2.4, 0.55,
                               edgecolor=color, facecolor=color+"22",
                               linewidth=1.5, zorder=3)
        ax.add_patch(ell)
        ax.text(x, y, text, ha="center", va="center",
                fontsize=7.5, color="#111", zorder=4)

    # Actors
    actor(0.5, 6.2, "Guest\nUser", GRAY)
    actor(0.5, 3.6, "Student", BLUE)
    actor(10.4, 5.0, "Instructor", GREEN)

    # Guest use cases
    for i, uc in enumerate(["Browse Courses", "View Course Details",
                             "Preview Free Lecture", "Register / Login"]):
        usecase(3.8, 6.9 - i*0.75, uc, GRAY)
        ax.annotate("", xy=(2.6, 6.9 - i*0.75), xytext=(0.85, 6.4),
                    arrowprops=dict(arrowstyle="-", color=GRAY, lw=0.9))

    # Student use cases (left column)
    student_ucs = ["Enroll in Course", "Purchase Course (PayPal)",
                   "Watch Video Lectures", "Track Progress",
                   "Chat with AI Tutor", "Take Quiz", "View Certificate"]
    for i, uc in enumerate(student_ucs):
        usecase(3.8, 5.0 - i*0.72, uc, BLUE)
        ax.annotate("", xy=(2.6, 5.0 - i*0.72), xytext=(0.85, 3.8),
                    arrowprops=dict(arrowstyle="-", color=BLUE, lw=0.9))

    # Instructor use cases (right column)
    instr_ucs = ["Create / Edit Course", "Upload Video Lectures",
                 "Generate AI Outline", "Generate AI Quiz",
                 "Publish / Unpublish", "View Enrolled Students"]
    for i, uc in enumerate(instr_ucs):
        usecase(7.2, 6.5 - i*0.75, uc, GREEN)
        ax.annotate("", xy=(8.4, 6.5 - i*0.75), xytext=(9.95, 5.2),
                    arrowprops=dict(arrowstyle="-", color=GREEN, lw=0.9))

    ax.set_title("Figure 3.4 — Use Case Diagram (SkillSphere LMS)",
                 fontsize=11, fontweight="bold", pad=10)
    return _save(fig)


def diag_deployment():
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6.5); ax.axis("off")
    fig.patch.set_facecolor("white")

    def box(x, y, w, h, color, title, lines):
        r = FancyBboxPatch((x, y), w, h,
                           boxstyle="round,pad=0.08",
                           edgecolor=color, facecolor=color + "22", lw=2)
        ax.add_patch(r)
        ax.text(x+w/2, y+h-0.22, title,
                ha="center", va="top", fontsize=9, fontweight="bold", color=color)
        for i, l in enumerate(lines):
            ax.text(x+w/2, y+h-0.5-i*0.28, l,
                    ha="center", fontsize=7.5, color="#333")

    # Browser
    box(0.2, 2.8, 2.0, 1.4, BLUE, "Browser", ["React 18 SPA", "Axios + SSE", "sessionStorage JWT"])

    # Vercel
    big = FancyBboxPatch((2.6, 1.8), 5.4, 3.6,
                         boxstyle="round,pad=0.1",
                         edgecolor=LBLUE, facecolor="#EEF4FB", lw=2)
    ax.add_patch(big)
    ax.text(5.3, 5.25, "Vercel Cloud",
            ha="center", fontsize=10, fontweight="bold", color=LBLUE)

    box(2.8, 2.9, 2.3, 2.1, LBLUE, "Edge CDN",
        ["React Build", "HTML / JS / CSS", "Global Cache"])
    box(5.5, 2.9, 2.3, 2.1, PURPLE, "Serverless Fn",
        ["Express.js API", "JWT Middleware", "Cold-start cache"])

    # External services
    ext = [
        (0.2, 0.2, BLUE,  "MongoDB Atlas",  ["M0 Cluster", "7 Collections", "Connection Pool"]),
        (3.0, 0.2, ORANGE,"Cloudinary CDN",  ["Video Storage", "Auto Transcoding", "Public IDs"]),
        (5.8, 0.2, GREEN, "Google Gemini",   ["gemini-2.0-flash", "JSON + SSE Mode", "Free Tier"]),
        (8.6, 0.2, "#C0392B","PayPal API",   ["Sandbox Mode", "Order Create", "Payment Capture"]),
    ]
    for x, y, c, t, ls in ext:
        box(x, y, 2.5, 1.5, c, t, ls)

    # Arrows
    ax.annotate("", xy=(2.6, 3.5), xytext=(2.2, 3.5),
                arrowprops=dict(arrowstyle="<->", color=GRAY, lw=1.8))
    ax.text(2.4, 3.7, "HTTPS", fontsize=7, color=GRAY, ha="center")

    ax.annotate("", xy=(5.5, 3.5), xytext=(5.1, 3.5),
                arrowprops=dict(arrowstyle="<->", color=GRAY, lw=1.5))

    for xi in [1.45, 4.25, 7.05, 9.85]:
        ax.annotate("", xy=(xi, 1.75), xytext=(min(xi, 7.8), 2.9),
                    arrowprops=dict(arrowstyle="-|>",
                                   color=GRAY, lw=1.2, linestyle="dashed"))

    ax.set_title("Figure 3.5 — Deployment Architecture Diagram",
                 fontsize=11, fontweight="bold", pad=10)
    return _save(fig)


def diag_sequence_payment():
    """Simple sequence diagram for PayPal payment flow."""
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10); ax.set_ylim(0, 7); ax.axis("off")
    fig.patch.set_facecolor("white")

    actors = ["Student\nBrowser", "React\nApp", "Express\nServer", "PayPal\nAPI", "MongoDB\nAtlas"]
    xs = [1, 2.7, 4.8, 7.0, 9.0]
    colors = [BLUE, LBLUE, PURPLE, ORANGE, GREEN]

    for x, a, c in zip(xs, actors, colors):
        ax.text(x, 6.7, a, ha="center", va="top", fontsize=9,
                fontweight="bold", color=c)
        ax.plot([x, x], [6.4, 0.3], color=c, lw=1.2, linestyle="dashed", alpha=0.6)

    messages = [
        (1, 2.7, 6.1, "Click 'Purchase'", False),
        (2.7, 4.8, 5.7, "POST /order/create", False),
        (4.8, 7.0, 5.3, "payment.create()", False),
        (7.0, 4.8, 4.9, "approvalUrl", True),
        (4.8, 7.0, 4.5, "Order.save(pending)", False),
        (4.8, 2.7, 4.1, "{approveUrl}", True),
        (2.7, 1,   3.7, "redirect to PayPal", False),
        (1,   7.0, 3.3, "PayPal Checkout UI", False),
        (7.0, 1,   2.9, "/payment-return?paymentId", True),
        (1,   4.8, 2.5, "POST /order/capture", False),
        (4.8, 9.0, 2.1, "Update Order + enroll student", False),
        (4.8, 2.7, 1.7, "{success}", True),
        (2.7, 1,   1.3, "Redirect to course", True),
    ]

    for x1, x2, y, msg, ret in messages:
        col = "#B03A2E" if ret else "#1A5276"
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="-|>", color=col, lw=1.5,
                                   linestyle="--" if ret else "-"))
        mx = (x1 + x2) / 2
        ax.text(mx, y + 0.12, msg, ha="center", fontsize=7.2, color=col)

    ax.set_title("Figure 3.6 — Sequence Diagram: PayPal Course Purchase Flow",
                 fontsize=11, fontweight="bold", pad=10)
    return _save(fig)


def diag_progress_flow():
    fig, ax = plt.subplots(figsize=(9, 7))
    ax.set_xlim(0, 9); ax.set_ylim(0, 7); ax.axis("off")
    fig.patch.set_facecolor("white")

    def node(x, y, w, h, color, text, shape="rect"):
        if shape == "diamond":
            dx, dy = w/2, h/2
            poly = plt.Polygon(
                [(x, y+dy), (x+dx, y+2*dy), (x+2*dx, y+dy), (x+dx, y)],
                edgecolor=color, facecolor=color+"33", lw=2, zorder=3)
            ax.add_patch(poly)
            ax.text(x+dx, y+dy, text, ha="center", va="center",
                    fontsize=8, color="#111", zorder=4)
        elif shape == "oval":
            e = mpatches.Ellipse((x+w/2, y+h/2), w, h,
                                 edgecolor=color, facecolor=color+"33",
                                 lw=2, zorder=3)
            ax.add_patch(e)
            ax.text(x+w/2, y+h/2, text, ha="center", va="center",
                    fontsize=8.5, fontweight="bold", color=color, zorder=4)
        else:
            r = FancyBboxPatch((x, y), w, h,
                               boxstyle="round,pad=0.08",
                               edgecolor=color, facecolor=color+"22",
                               lw=1.8, zorder=3)
            ax.add_patch(r)
            ax.text(x+w/2, y+h/2, text, ha="center", va="center",
                    fontsize=8, color="#111", zorder=4)
        return (x+w/2, y+h/2)

    c = node(3.3, 6.2, 2.4, 0.55, BLUE, "Student opens course", shape="oval")
    n1 = node(3.3, 5.2, 2.4, 0.65, BLUE, "Watch Lecture")
    n2 = node(3.3, 4.3, 2.4, 0.65, BLUE, "Mark as Viewed\n(POST /mark-lecture-viewed)")
    n3 = node(2.0, 3.2, 1.8, 0.65, LBLUE, "Progress % Updated\n(server-side calc)")
    n4 = node(4.5, 3.1, 1.9, 0.8, ORANGE, "All lectures\nviewed?", shape="diamond")
    n5 = node(3.3, 1.9, 2.4, 0.65, GREEN, "Mark course Completed\n+ set completionDate")
    n6 = node(3.3, 1.0, 2.4, 0.65, GREEN, "Certificate Generated")

    pairs = [(c, n1, ""), (n1, n2, ""), (n2, n3, ""),
             (n2, n4, ""), (n3, n4, "sync"),
             (n4, n5, "Yes"), (n5, n6, "")]

    for (x1, y1), (x2, y2), lbl in pairs:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=GRAY, lw=1.5))
        if lbl:
            ax.text((x1+x2)/2 + 0.2, (y1+y2)/2, lbl,
                    fontsize=8, color=GRAY)

    # No arrow back
    ax.annotate("", xy=(n1[0]-1.5, n1[1]), xytext=(n4[0]+1, n4[1]),
                arrowprops=dict(arrowstyle="-|>", color=ORANGE, lw=1.3,
                                connectionstyle="arc3,rad=-0.5"))
    ax.text(7.3, 3.5, "No — next\nlecture", fontsize=8, color=ORANGE, ha="center")

    ax.set_title("Figure 3.7 — Activity Diagram: Student Progress Tracking Flow",
                 fontsize=11, fontweight="bold", pad=10)
    return _save(fig)


# ─────────────────────────────────────────────────────────────
# TABLE HELPERS
# ─────────────────────────────────────────────────────────────

def add_2col_table(doc, headers, rows, col_widths=(2.5, 4.0)):
    spacer(doc, pts=8)           # clear space before every table
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    hrow = table.rows[0]
    shade_row(hrow, "2C5F8A")
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        drow = table.rows[ri + 1]
        if ri % 2 == 0:
            shade_row(drow, "EBF0F7")
        for ci, val in enumerate(row_data):
            c = drow.cells[ci]
            run = c.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # set widths
    for row in table.rows:
        for ci, w in enumerate(col_widths):
            row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────────────────────
# THESIS ASSEMBLY
# ─────────────────────────────────────────────────────────────

def build_thesis():
    doc = Document()
    set_margins(doc, top=1, bottom=1, left=1.25, right=1)

    # ── PAGE 1: Title ──────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SkillSphere: Design and Development of an\nAI-Powered Learning Management System")
    set_font(r, size=16, bold=True)
    para_spacing(p, before=30, after=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A thesis submitted in partial fulfillment of the requirements for\nthe award of the degree of")
    set_font(r, size=12, italic=True)
    para_spacing(p, before=6, after=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BACHELOR OF TECHNOLOGY\nin\nComputer Engineering")
    set_font(r, size=14, bold=True)
    para_spacing(p, before=6, after=12)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Submitted by").font.size = Pt(12)
    para_spacing(p, before=6, after=6)

    for name, reg in [("Aditya Pandey", "2021UG----"), ("<Team Member 2>", "2021UG----")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{name}\n({reg})")
        set_font(r, size=13, bold=True)
        para_spacing(p, before=2, after=2)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Under the Supervision of").font.size = Pt(12)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("<Supervisor Name>")
    set_font(r, size=13, bold=True)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(Department of Electronics & Computer Engineering)")
    set_font(r, size=12, bold=True)
    para_spacing(p, before=4, after=18)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("National Institute of Advanced Manufacturing Technology\n"
                  "Deemed to be University (Distinct Category)\n"
                  "Hatia, Ranchi – 834 003 (Jharkhand)\n"
                  "January 2026 – May 2026")
    set_font(r, size=12, bold=True)
    para_spacing(p, before=6, after=6)

    doc.add_page_break()

    # ── Candidate Declaration ───────────────────────────────────
    heading(doc, "CANDIDATE DECLARATION")
    dec_text = (
        "I hereby certify that I have properly checked and verified all items as prescribed "
        "in the checklist and ensure that this thesis is in the proper format specified in "
        "the guidelines for thesis preparation.\n\n"
        "I also declare that the work contained in this report is my own. I understand that "
        "plagiarism is defined as any one or combination of the following:\n"
        "1. To steal and pass off the ideas or words of another as one's own.\n"
        "2. To use another's production without crediting the source.\n"
        "3. To commit literary theft.\n"
        "4. To present as new and original an idea derived from an existing source.\n\n"
        "I affirm that no portion of my work is plagiarized, and the experiments and results "
        "reported in this thesis are not manipulated. In the event of any complaint of plagiarism "
        "or manipulation of results, I shall be fully responsible and answerable. My faculty "
        "supervisor(s) will not be responsible for the same."
    )
    body(doc, dec_text)
    doc.add_paragraph()
    for label in ["Signature:", "Name:", "Roll No.:", "Date:"]:
        p = doc.add_paragraph(); run = p.add_run(label)
        set_font(run, size=12)
    doc.add_page_break()

    # ── Thesis Certificate ──────────────────────────────────────
    heading(doc, "THESIS CERTIFICATE")
    cert = (
        'This is to certify that the thesis entitled "SkillSphere: Design and Development of '
        'an AI-Powered Learning Management System", submitted in partial fulfillment of the '
        'requirement for the award of the Degree of Bachelor of Technology in Computer Engineering '
        'by Aditya Pandey (Reg. No. 2021UG----) & <Team Member 2> (Reg. No. 2021UG----) to the '
        'National Institute of Advanced Manufacturing Technology, Ranchi, is an authentic record of '
        'the project work carried out during the period January 2026 to May 2026 under the '
        'supervision of <Supervisor Name>. The contents of this thesis, in full or in parts, have '
        'not been submitted to any other Institute or University for the award of any degree or '
        'diploma to the best of my knowledge.'
    )
    body(doc, cert)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Date:                                                            "
                    "Signature of the Supervisor(s)")
    set_font(run, size=12)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("                                                                               "
                    "External Examiner")
    set_font(run, size=12)
    doc.add_page_break()

    # ── Abstract ────────────────────────────────────────────────
    heading(doc, "ABSTRACT")
    abstract = (
        "The rapid growth of online education has created demand for learning platforms that do "
        "more than simply host recorded lectures. Yet even the most widely used systems still treat "
        "content creation and student support as entirely manual processes, leaving instructors to "
        "build curricula from scratch and learners to seek help outside the platform whenever they "
        "get stuck. This thesis presents SkillSphere, a full-stack Learning Management System "
        "built on the MERN stack (MongoDB, Express.js, React 18, Node.js) that integrates Google "
        "Gemini 2.0 Flash as a first-class architectural component rather than an optional add-on.\n\n"
        "SkillSphere addresses three concrete pain points. First, course creation is accelerated "
        "through AI-generated outlines — an instructor enters a topic and difficulty level, and "
        "Gemini returns a complete, structured curriculum in under seven seconds. Second, "
        "assessment quality is improved through AI-generated multiple-choice quizzes grouped by "
        "lecture topic, with configurable difficulty distributions across easy, medium, and hard "
        "tiers. Third, learner support is embedded directly in the platform through a real-time "
        "streaming AI tutor that is scoped to the specific course content, preventing the "
        "distraction of general-purpose search engines.\n\n"
        "The platform handles the full e-learning lifecycle: role-based authentication with JWT, "
        "video hosting via Cloudinary CDN, payment processing through PayPal REST SDK, lecture-level "
        "progress tracking with automatic completion detection, certificate generation, and "
        "unenrollment with data cleanup. The backend is deployed as Vercel serverless functions "
        "connected to MongoDB Atlas, with deliberate cold-start mitigations including connection "
        "caching and extended Axios timeouts. Functional testing across thirteen test cases confirms "
        "all core workflows operate correctly. Performance measurements show warm API responses in "
        "100–300 ms and AI generation in three to twelve seconds depending on complexity.\n\n"
        "The results demonstrate that treating generative AI as a core architectural layer — not "
        "a plugin — meaningfully reduces instructor workload and improves the quality of student "
        "support without requiring additional infrastructure."
    )
    body(doc, abstract)
    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    set_font(run, bold=True, size=12)
    krun = p.add_run("LMS, MERN Stack, Generative AI, Gemini 2.0 Flash, E-Learning, "
                     "JWT Authentication, PayPal Integration, Cloudinary CDN, "
                     "Server-Sent Events, MongoDB Atlas.")
    set_font(krun, size=12)
    doc.add_page_break()

    # ── Acknowledgement ─────────────────────────────────────────
    heading(doc, "ACKNOWLEDGEMENT")
    ack = (
        "I would like to express my sincere gratitude to the Director, National Institute of "
        "Advanced Manufacturing Technology (NIAMT), Ranchi, for providing the infrastructure and "
        "academic environment that made this project possible. I am deeply thankful to Dr. Madhu "
        "Kumari, Head of the Department, Electronics and Computer Engineering, for her timely "
        "guidance throughout the academic year.\n\n"
        "My heartfelt thanks go to my project supervisor, <Supervisor Name>, whose consistent "
        "feedback helped me refine both the technical architecture and the written documentation "
        "of this work. Every iteration of the system — from the initial proof-of-concept to the "
        "deployed Vercel application — benefited from the constructive questions raised during "
        "our review sessions.\n\n"
        "I would also like to thank my fellow students in the Electronics and Computer Engineering "
        "department whose discussions, however informal, often pushed me to think through design "
        "decisions more carefully. Finally, I am grateful to my family for their patience during "
        "the months spent debugging late into the night."
    )
    body(doc, ack)
    doc.add_page_break()

    # ── Table of Contents ───────────────────────────────────────
    heading(doc, "TABLE OF CONTENTS")
    toc_entries = [
        ("Abstract", "iii"),
        ("Acknowledgement", "iv"),
        ("List of Figures", "vi"),
        ("List of Abbreviations", "vii"),
        ("Chapter 1: Introduction", "1"),
        ("    1.1  Background and Motivation", "1"),
        ("    1.2  Problem Statement", "2"),
        ("    1.3  Objectives of the Project", "3"),
        ("    1.4  Outline of the Thesis", "3"),
        ("Chapter 2: Literature Review", "4"),
        ("    2.1  Overview of E-Learning Platforms", "4"),
        ("    2.2  AI in Education: Summary of Key Studies", "5"),
        ("    2.3  Identified Research Gaps", "6"),
        ("Chapter 3: Methodology and Design", "8"),
        ("    3.1  Technology Stack", "8"),
        ("    3.2  System Architecture", "9"),
        ("    3.3  Database Design", "12"),
        ("    3.4  UML Diagrams", "14"),
        ("    3.5  AI Integration Design", "16"),
        ("    3.6  Security Architecture", "18"),
        ("Chapter 4: Results and Discussion", "20"),
        ("    4.1  Functional Testing Results", "20"),
        ("    4.2  Performance Observations", "21"),
        ("    4.3  Discussion", "22"),
        ("Chapter 5: Conclusion", "24"),
        ("    5.1  Conclusion", "24"),
        ("    5.2  Limitations", "25"),
        ("    5.3  Future Scope", "25"),
        ("References", "27"),
    ]
    for title, page in toc_entries:
        p = doc.add_paragraph()
        is_chapter = title.startswith("Chapter") or title in ("Abstract","Acknowledgement","List of Figures","List of Abbreviations","References")
        run = p.add_run(title)
        set_font(run, size=11, bold=is_chapter)
        dots = "." * max(2, 65 - len(title))
        run2 = p.add_run(f"  {dots}  {page}")
        set_font(run2, size=11)
        para_spacing(p, before=1, after=1)
    doc.add_page_break()

    # ── List of Figures ─────────────────────────────────────────
    heading(doc, "LIST OF FIGURES")
    figs = [
        ("Fig 3.1", "SkillSphere Three-Tier Architecture", "9"),
        ("Fig 3.2", "Entity Relationship Diagram", "12"),
        ("Fig 3.3", "AI Feature Pipeline (Gemini 2.0 Flash Integration)", "16"),
        ("Fig 3.4", "Use Case Diagram", "14"),
        ("Fig 3.5", "Deployment Architecture Diagram", "18"),
        ("Fig 3.6", "Sequence Diagram: PayPal Course Purchase Flow", "15"),
        ("Fig 3.7", "Activity Diagram: Student Progress Tracking Flow", "17"),
    ]
    table = doc.add_table(rows=1 + len(figs), cols=3)
    table.style = "Table Grid"
    shade_row(table.rows[0], "2C5F8A")
    for i, h in enumerate(["Figure No.", "Figure Title", "Page No."]):
        c = table.rows[0].cells[i]
        r = c.paragraphs[0].add_run(h)
        set_font(r, size=11, bold=True, color=(255, 255, 255))
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, (fn, ft, pg) in enumerate(figs):
        row = table.rows[ri + 1]
        if ri % 2 == 0:
            shade_row(row, "EBF0F7")
        for ci, val in enumerate([fn, ft, pg]):
            c = row.cells[ci]
            r = c.paragraphs[0].add_run(val)
            set_font(r, size=11)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ── List of Abbreviations ───────────────────────────────────
    heading(doc, "LIST OF ABBREVIATIONS")
    abbrevs = [
        ("AI",    "Artificial Intelligence"),
        ("API",   "Application Programming Interface"),
        ("CDN",   "Content Delivery Network"),
        ("CORS",  "Cross-Origin Resource Sharing"),
        ("CRUD",  "Create, Read, Update, Delete"),
        ("CSS",   "Cascading Style Sheets"),
        ("DB",    "Database"),
        ("ERD",   "Entity Relationship Diagram"),
        ("HLD",   "High Level Design"),
        ("HTTP",  "Hypertext Transfer Protocol"),
        ("JWT",   "JSON Web Token"),
        ("LLD",   "Low Level Design"),
        ("LMS",   "Learning Management System"),
        ("MCQ",   "Multiple Choice Question"),
        ("MERN",  "MongoDB, Express.js, React, Node.js"),
        ("ODM",   "Object Document Mapper"),
        ("REST",  "Representational State Transfer"),
        ("SDK",   "Software Development Kit"),
        ("SPA",   "Single Page Application"),
        ("SQL",   "Structured Query Language"),
        ("SSE",   "Server-Sent Events"),
        ("UI",    "User Interface"),
        ("UML",   "Unified Modeling Language"),
        ("URL",   "Uniform Resource Locator"),
        ("UX",    "User Experience"),
    ]
    table = doc.add_table(rows=1 + len(abbrevs), cols=2)
    table.style = "Table Grid"
    shade_row(table.rows[0], "2C5F8A")
    for i, h in enumerate(["Abbreviation", "Full Form"]):
        c = table.rows[0].cells[i]
        r = c.paragraphs[0].add_run(h)
        set_font(r, size=11, bold=True, color=(255, 255, 255))
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, (ab, full) in enumerate(abbrevs):
        row = table.rows[ri + 1]
        if ri % 2 == 0:
            shade_row(row, "EBF0F7")
        row.cells[0].paragraphs[0].add_run(ab).font.size = Pt(11)
        row.cells[1].paragraphs[0].add_run(full).font.size = Pt(11)
    doc.add_page_break()

    # ────────────────────────────────────────────────────────────
    # CHAPTER 1: INTRODUCTION
    # ────────────────────────────────────────────────────────────
    heading(doc, "CHAPTER 1: INTRODUCTION", level=1)

    section_heading(doc, "1.1  Background and Motivation")
    body(doc,
         "The global e-learning market was valued at USD 250 billion in 2023 and is projected "
         "to reach approximately USD 1 trillion by 2032 (Global Market Insights, 2023). The "
         "COVID-19 pandemic forced educational institutions worldwide to shift online, and even "
         "after physical campuses reopened, the appetite for flexible, self-paced learning "
         "remained. Platforms such as Coursera, Udemy, and edX reported enrollment growth "
         "between 200% and 500% during 2020–2022, and that momentum has largely held.")
    body(doc,
         "Yet, despite the explosion in learner numbers, the fundamental experience on most "
         "platforms has changed very little. Instructors still spend days or weeks designing "
         "course structures, writing learning objectives, and crafting quiz questions — tasks "
         "that require instructional design expertise that many domain-expert instructors simply "
         "do not have. Students, meanwhile, still watch videos passively and are left to search "
         "Google or Stack Overflow when a concept does not click, breaking their flow and "
         "diluting the learning experience the instructor intended.")
    body(doc,
         "The emergence of large language models — particularly Google Gemini and OpenAI's GPT "
         "series — has opened a genuine opportunity to address these gaps. Tools like ChatGPT "
         "have already changed how many people write, code, and research. What has been slower "
         "to arrive is their integration into learning platforms at the architectural level, "
         "not merely as a chat widget bolted onto an otherwise unchanged system. The motivation "
         "behind SkillSphere is precisely this gap: to build an LMS where AI is a structural "
         "component, not an afterthought.")

    section_heading(doc, "1.2  Problem Statement")
    body(doc,
         "Existing Learning Management Systems suffer from three persistent shortcomings "
         "when viewed through the lens of modern generative AI capabilities:")
    for bullet in [
        "Instructor overhead: Creating a well-structured course curriculum with appropriate "
        "learning objectives, lecture sequences, and assessments remains a largely manual, "
        "time-intensive process. Instructors who are experts in their subject matter are not "
        "necessarily trained instructional designers, and current platforms provide no "
        "intelligent assistance during content creation.",
        "Passive learning experience: Once enrolled, students watch recorded lectures with "
        "no in-platform mechanism for asking follow-up questions tied to the specific course "
        "content. They must either wait for a forum response (which may arrive hours later) "
        "or search externally — both of which disrupt the learning context.",
        "Static, manually created assessments: Quiz questions are authored once by the "
        "instructor and remain unchanged regardless of which lectures a student has completed "
        "or how they have performed so far. There is no mechanism for difficulty calibration "
        "or lecture-specific question grouping."
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bullet)
        set_font(run, size=12)
        para_spacing(p, before=2, after=4)
    body(doc,
         "SkillSphere addresses all three problems by embedding Google Gemini 2.0 Flash "
         "into three distinct points in the course lifecycle: outline generation at creation "
         "time, quiz generation after curriculum is defined, and real-time tutoring during "
         "active learning.")

    section_heading(doc, "1.3  Objectives of the Project")
    objectives = [
        "Design and implement a full-stack LMS using the MERN stack (MongoDB, Express.js, React 18, Node.js).",
        "Integrate Google Gemini 2.0 Flash API for AI-assisted course outline generation, adaptive quiz creation, and real-time streaming tutoring.",
        "Implement secure role-based authentication using JSON Web Tokens with Instructor and Student roles.",
        "Develop a complete e-commerce flow for paid courses using the PayPal REST SDK.",
        "Build a media management pipeline using Cloudinary CDN for instructor video uploads.",
        "Implement real-time AI chat using Server-Sent Events for token-by-token streaming.",
        "Create lecture-level progress tracking with automatic completion detection and certificate generation.",
        "Deploy a production-ready application on Vercel with MongoDB Atlas, incorporating serverless cold-start mitigations.",
        "Demonstrate the feasibility of an AI-native LMS architecture as an improvement over AI-as-add-on designs."
    ]
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {obj}")
        set_font(run, size=12)
        para_spacing(p, before=2, after=3)

    section_heading(doc, "1.4  Outline of the Thesis")
    body(doc,
         "Chapter 2 reviews related work in e-learning platforms and AI applications in "
         "education, identifying the gaps that SkillSphere addresses. Chapter 3 presents the "
         "full methodology and system design, including technology selection, architecture "
         "diagrams, database schema, UML models, and the AI integration design. Chapter 4 "
         "reports the results of functional testing and performance measurements, with "
         "discussion of what the results mean in practice. Chapter 5 concludes with a summary "
         "of contributions, known limitations, and directions for future work.")
    doc.add_page_break()

    # ────────────────────────────────────────────────────────────
    # CHAPTER 2: LITERATURE REVIEW
    # ────────────────────────────────────────────────────────────
    heading(doc, "CHAPTER 2: LITERATURE REVIEW", level=1)

    section_heading(doc, "2.1  Overview of E-Learning Platforms")
    body(doc,
         "Online learning platforms have evolved through three broad generations. The first "
         "generation, represented by early SCORM-based systems such as Moodle, focused on "
         "packaging and delivering existing instructional material. The second generation, "
         "exemplified by MOOCs like Coursera and edX launched around 2012, scaled that delivery "
         "to hundreds of thousands of learners at once. The third and current generation has "
         "added commerce, community features, and increasingly, AI-powered recommendations. "
         "Oblinger (2012) argued that the most impactful educational technology would be that "
         "which adapts the learning environment to the individual learner rather than expecting "
         "the learner to adapt to the technology. This framing predated large language models "
         "but anticipated them with remarkable accuracy.")
    body(doc,
         "Siemens (2005) introduced connectivism as a learning theory suited to the digital "
         "age, emphasising that learning occurs across a network of connections rather than "
         "solely inside a single learner's head. An AI tutor embedded in a course platform "
         "is a natural realisation of this idea — it represents an always-available node in the "
         "learner's network, one that is specifically calibrated to the course content rather "
         "than general knowledge.")
    body(doc,
         "From a platform design perspective, Alario-Hoyos et al. (2017) studied deployment "
         "patterns for MOOCs and concluded that personalisable assessment and timely feedback "
         "were the two features most consistently correlated with course completion. SkillSphere "
         "addresses both: AI-generated quizzes with per-question explanations provide immediate "
         "feedback, while the streaming tutor provides on-demand clarification without requiring "
         "instructor availability.")

    section_heading(doc, "2.2  AI in Education: Summary of Key Studies")
    body(doc,
         "The application of AI to education has a research history going back to intelligent "
         "tutoring systems in the 1980s, but the arrival of transformer-based language models "
         "has sharply changed what is practically achievable. Bhatt and Mudge (2021) conducted "
         "a systematic review of machine learning applications in education and found that most "
         "deployed systems fell into three categories: automated essay grading, early warning "
         "systems for at-risk students, and recommendation engines for course sequencing. None "
         "of the reviewed systems integrated AI into the content creation process itself.")
    body(doc,
         "Pinkwart (2016) identified the gap between research prototypes and deployed systems "
         "as one of the central challenges for the field, noting that most intelligent tutoring "
         "systems required specialist knowledge engineering effort that made them impractical "
         "to scale. Large language models fundamentally change this equation: Gemini and GPT-4 "
         "require prompt engineering rather than knowledge engineering, making them accessible "
         "to developers without deep expertise in educational psychology.")
    body(doc,
         "Vaswani et al. (2017) introduced the transformer architecture that underlies all "
         "modern large language models, including the Gemini family used in SkillSphere. "
         "Brown et al. (2020) demonstrated that models of sufficient scale could perform "
         "few-shot learning — generating structured outputs from brief natural language "
         "descriptions — without task-specific fine-tuning. This capability is exactly what "
         "SkillSphere exploits: the model receives a course topic and a JSON schema in the "
         "prompt and returns a complete, schema-compliant curriculum without any "
         "domain-specific training.")
    body(doc,
         "Meza-Kubo and Morán (2020) systematically reviewed technologies for online learning "
         "and noted that platforms combining synchronous and asynchronous interaction produced "
         "better engagement metrics than purely asynchronous designs. The AI tutor in SkillSphere "
         "simulates synchronous availability at asynchronous cost: it is always ready, "
         "contextually aware, and responds in real time via streaming.")

    section_heading(doc, "2.3  Identified Research Gaps")
    body(doc, "Based on the literature review, five clear gaps emerge:")
    gaps = [
        ("Gap 1 — AI-native content creation:",
         "Most LMS research addresses AI for learner analytics or recommendations. "
         "Using AI to assist instructors during course creation has not been widely "
         "studied or deployed at scale."),
        ("Gap 2 — Integrated AI toolchain:",
         "Where AI tools for instructors do exist, they are typically disconnected from "
         "the LMS — an instructor uses ChatGPT externally and pastes the result into "
         "the platform. There is no seamless pipeline from AI generation to course "
         "publication within a single system."),
        ("Gap 3 — Course-scoped AI tutor:",
         "Chatbot integrations in LMS platforms typically connect to a general-purpose "
         "model with no awareness of the specific course content. A tutor scoped to one "
         "course's objectives and curriculum is rare in open, deployable systems."),
        ("Gap 4 — Lecture-grouped adaptive assessment:",
         "AI-generated quizzes calibrated to specific lecture subsets with configurable "
         "difficulty distributions have not been documented as a production feature in "
         "open LMS platforms."),
        ("Gap 5 — Full-stack MERN + Gemini implementation:",
         "Published academic literature contains very few complete implementation studies "
         "of Gemini API integration in a full-stack web application, particularly in the "
         "context of education.")
    ]
    for title, detail in gaps:
        p = doc.add_paragraph()
        r1 = p.add_run(title + " ")
        set_font(r1, size=12, bold=True)
        r2 = p.add_run(detail)
        set_font(r2, size=12)
        para_spacing(p, before=4, after=4)

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────
    # CHAPTER 3: METHODOLOGY AND DESIGN
    # ────────────────────────────────────────────────────────────
    heading(doc, "CHAPTER 3: METHODOLOGY AND DESIGN", level=1)

    section_heading(doc, "3.1  Technology Stack")
    body(doc,
         "SkillSphere is built entirely in JavaScript/TypeScript-compatible JavaScript, "
         "sharing the same language across client and server. This choice eliminates "
         "context switching between languages and allows data models defined once to be "
         "consumed without transformation on either side. The technology selections below "
         "were each driven by a specific constraint or advantage.")

    section_heading(doc, "3.1.1  Frontend Technologies", size=11)
    add_2col_table(doc,
        ["Technology (Version)", "Role and Justification"],
        [
            ("React 18 + Vite 5", "Core UI library. Vite provides near-instant HMR. All pages lazy-loaded via React.lazy() + Suspense to minimise initial bundle size."),
            ("React Router DOM 6", "Client-side routing with nested routes and a RouteGuard component for role-based access control."),
            ("Tailwind CSS 3 + Radix UI", "Utility-first styling with accessible headless primitives. cn() from lib/utils.js merges conditional class strings safely."),
            ("Framer Motion 11", "Declarative page transition and component mount animations without imperative DOM manipulation."),
            ("Axios 1 + axiosInstance", "HTTP client with a request interceptor that automatically injects the JWT from sessionStorage and enforces the 12-second timeout."),
            ("React Player 2", "Handles both Cloudinary-hosted video (mp4) and YouTube-embedded lectures in the same component."),
        ], col_widths=(2.0, 4.5))

    section_heading(doc, "3.1.2  Backend Technologies", size=11)
    add_2col_table(doc,
        ["Technology (Version)", "Role and Justification"],
        [
            ("Node.js 18 + Express 4", "REST API server. CommonJS module system throughout server code, consistent with the Vercel serverless runtime."),
            ("Mongoose 7", "MongoDB ODM providing schema validation, type casting, and the array filter syntax used for nested lecture updates."),
            ("jsonwebtoken 9 + bcryptjs 2", "Stateless JWT auth (120-minute expiry). bcrypt with 10 salt rounds adds ~100 ms per hash, intentionally slowing brute-force attempts."),
            ("@google/generative-ai", "Google's official Gemini SDK. Used in JSON mode for outlines and quizzes; streaming mode for the AI tutor chat."),
            ("Cloudinary SDK 2", "Handles upload, deletion, and duration metadata queries for instructor video content."),
            ("PayPal REST SDK 1", "Manages the create-order → redirect → capture payment lifecycle in sandbox mode."),
            ("Multer 1", "Disk-storage middleware that stages multipart video uploads to server/uploads/ before forwarding to Cloudinary."),
        ], col_widths=(2.0, 4.5))

    section_heading(doc, "3.2  System Architecture")
    body(doc,
         "SkillSphere follows a classic three-tier architecture adapted for cloud-native, "
         "serverless deployment on Vercel. The presentation tier is a React 18 Single Page "
         "Application. The application tier is an Express.js REST API deployed as Vercel "
         "Serverless Functions. The data tier combines MongoDB Atlas for structured document "
         "storage and Cloudinary CDN for binary media. Three external services — Google Gemini "
         "API, PayPal Sandbox, and Cloudinary — are accessed exclusively from the application "
         "tier, keeping credentials server-side and never exposed to the browser.")

    # Architecture diagram
    buf = diag_three_tier()
    insert_figure(doc, buf, 5.8, "Figure 3.1 — SkillSphere Three-Tier Architecture")

    section_heading(doc, "3.2.1  Logical Service Boundaries", size=11)
    body(doc,
         "Although SkillSphere runs as a single Express application, the codebase is divided "
         "into eight logical service domains, each with its own route prefix, controller "
         "module, and clearly defined responsibility:")
    add_2col_table(doc,
        ["Route Prefix", "Responsibility"],
        [
            ("/auth",                    "Register, login, and JWT verification"),
            ("/instructor/course",       "Course CRUD for instructors"),
            ("/media",                   "Upload, delete, and bulk-upload to Cloudinary"),
            ("/student/course",          "Course listing, filtering, and detail view"),
            ("/student/order",           "PayPal order creation, capture, and free enrollment"),
            ("/student/courses-bought",  "Purchased course list and unenrollment"),
            ("/student/course-progress", "Lecture progress marking and completion tracking"),
            ("/ai",                      "Outline generation, quiz generation, attempts, and chat"),
        ], col_widths=(2.2, 4.3))

    section_heading(doc, "3.3  Database Design")
    body(doc,
         "SkillSphere uses MongoDB as its primary data store, accessed through Mongoose. "
         "The choice of a document store over a relational database reflects the hierarchical "
         "structure of course data: a Course document nests an array of Lecture sub-documents, "
         "and each Lecture nests metadata fields. This avoids JOIN operations that would be "
         "required in a relational model. Seven collections are defined.")

    # ER Diagram — own page so it doesn't crowd the text above
    doc.add_page_break()
    buf2 = diag_er()
    insert_figure(doc, buf2, 5.7, "Figure 3.2 — Entity Relationship Diagram (SkillSphere Database)")

    body(doc,
         "The User document stores credentials and role. The Course document is the central "
         "entity: it contains the full curriculum as a nested array of LectureSchema objects, "
         "each holding a Cloudinary URL and public identifier. StudentCourses tracks which "
         "courses a user has purchased, while CourseProgress tracks which lectures they have "
         "viewed. The Quiz collection stores one AI-generated quiz per course, indexed with a "
         "unique constraint on courseId. QuizAttempt records each student's submission with a "
         "compound unique index on (userId, courseId, groupIndex) — retaking a quiz simply "
         "overwrites the existing record rather than creating a new one.")

    # ── Section 3.4: UML Diagrams ── each diagram gets its own page ──
    doc.add_page_break()
    section_heading(doc, "3.4  UML Diagrams")
    body(doc,
         "The following diagrams document the system's actors and interactions using standard "
         "UML notation. Figure 3.4 shows the full use case scope. Figures 3.6 and 3.7 detail "
         "two key flows — payment processing and progress tracking — as sequence and activity "
         "diagrams respectively.")

    section_heading(doc, "3.4.1  Use Case Diagram", size=11)
    buf3 = diag_usecase()
    insert_figure(doc, buf3, 5.7, "Figure 3.4 — Use Case Diagram (SkillSphere LMS)")

    doc.add_page_break()
    section_heading(doc, "3.4.2  Sequence Diagram — PayPal Payment Flow", size=11)
    body(doc,
         "The PayPal payment flow spans three systems: the browser, the Express server, and "
         "PayPal's REST API. The sequence below shows the create-order, redirect, and capture "
         "stages, including the MongoDB writes that confirm enrollment.")
    buf4 = diag_sequence_payment()
    insert_figure(doc, buf4, 5.7, "Figure 3.6 — Sequence Diagram: PayPal Course Purchase Flow")

    doc.add_page_break()
    section_heading(doc, "3.4.3  Activity Diagram — Progress Tracking", size=11)
    body(doc,
         "Student progress is tracked at the individual lecture level. The activity diagram "
         "below shows how marking a lecture as viewed triggers a server-side percentage "
         "calculation and, once all lectures are viewed, sets the course completion flag and "
         "records the completion date.")
    buf5 = diag_progress_flow()
    insert_figure(doc, buf5, 4.8, "Figure 3.7 — Activity Diagram: Student Progress Tracking Flow")

    # ── Section 3.5: AI Integration ──
    doc.add_page_break()
    section_heading(doc, "3.5  AI Integration Design")
    body(doc,
         "The AI layer is built around a single Gemini helper module "
         "(server/helpers/gemini.js) that exposes a getModel() factory function. The factory "
         "accepts three parameters: a system instruction string, a maximum token limit, and a "
         "boolean flag that switches between JSON mode and streaming mode. "
         "Setting responseMimeType to 'application/json' instructs Gemini to return "
         "syntactically valid JSON without markdown code fences, which eliminates the parsing "
         "fragility that comes from stripping code block wrappers off model output.")

    # AI Pipeline diagram
    buf6 = diag_ai_pipeline()
    insert_figure(doc, buf6, 5.8, "Figure 3.3 — AI Feature Pipeline (Gemini 2.0 Flash Integration)")

    section_heading(doc, "3.5.1  Course Outline Generation", size=11)
    body(doc,
         "When an instructor clicks 'Generate with AI', the client sends a POST to "
         "/ai/generate-outline with the topic, difficulty level, target audience, and an "
         "optional syllabus hint. The server builds a prompt that includes the course "
         "parameters and a full JSON schema specifying the expected output structure: title, "
         "subtitle, description, objectives array, welcome message, and a curriculum of "
         "sections each containing an array of lectures. The model returns the schema-compliant "
         "JSON in under seven seconds on average. The server flattens the nested "
         "section-lecture hierarchy into a single lecture array for direct population "
         "of the curriculum form, and converts the objectives array to a newline-separated "
         "string matching the database field format.")

    section_heading(doc, "3.5.2  AI Quiz Generation", size=11)
    body(doc,
         "Quiz generation is triggered after the curriculum is finalised. The instructor "
         "configures a mode (end-of-course or interval-based), question count, and difficulty "
         "distribution. The server groups lectures according to the chosen mode and sends each "
         "group to Gemini with the lecture titles as context. The model generates the "
         "requested number of questions per group, each with four options, a correct answer "
         "index, a difficulty label, and an explanation. Explanations are returned to students "
         "after submission to provide corrective feedback regardless of whether their answer "
         "was right or wrong.")

    section_heading(doc, "3.5.3  AI Tutor Chat (Streaming SSE)", size=11)
    body(doc,
         "The AI tutor uses a fundamentally different response mode. Instead of waiting for "
         "a complete JSON response, the server calls sendMessageStream() on the Gemini chat "
         "object and writes each token to the HTTP response as an SSE event the moment it "
         "arrives. The response header Content-Type is set to text/event-stream, and each "
         "event is formatted as data: {\"chunk\": \"<token>\"} with a double newline terminator. "
         "The client reads the response body as a ReadableStream and calls a callback for "
         "each parsed chunk, appending text to the chat bubble in real time. "
         "A final data: {\"done\": true} event signals stream completion. Up to ten turns of "
         "conversation history are passed to model.startChat() on each request, giving the "
         "tutor conversational continuity within a session. The system instruction scopes "
         "the tutor to the specific course by including the course title, description, and "
         "learning objectives, and instructs the model to redirect off-topic questions "
         "back to course content.")

    # ── Section 3.6: Security ──
    doc.add_page_break()
    section_heading(doc, "3.6  Security Architecture")
    body(doc,
         "Authentication is stateless. On login, the server signs a JWT containing "
         "the user's ID, email, username, and role with a secret stored as an environment "
         "variable. The token expires after 120 minutes. The client stores it in "
         "sessionStorage rather than localStorage: sessionStorage clears when the browser tab "
         "closes, shortening the window during which a stolen token remains valid. An Axios "
         "request interceptor injects the token as a Bearer header on every API call, and the "
         "server's auth middleware verifies it with jwt.verify() before any protected "
         "handler runs.")
    body(doc,
         "Passwords are hashed with bcryptjs using 10 salt rounds, which takes approximately "
         "100 ms per hash on modern hardware. That delay is slow enough to make brute-force "
         "credential attacks impractical without any noticeable impact on legitimate login "
         "latency. CORS is configured with an explicit origin whitelist. NoSQL injection is "
         "mitigated by Mongoose's schema-based type casting, which rejects unexpected "
         "query operators at the ODM layer.")

    add_2col_table(doc,
        ["Security Concern", "Mitigation in SkillSphere"],
        [
            ("SQL / NoSQL Injection",     "Mongoose schema validation rejects malformed input"),
            ("Cross-Site Scripting (XSS)", "sessionStorage cleared on tab close; no innerHTML usage"),
            ("CSRF",                      "Stateless JWT — no session cookies to forge"),
            ("Brute-Force Login",         "bcrypt 10-round delay (~100 ms per attempt)"),
            ("Secret Exposure",           "dotenv locally; Vercel environment variables in production"),
            ("CORS Exploitation",         "Explicit origin whitelist on all API routes"),
            ("Unauthorised API Access",   "JWT middleware on all sensitive endpoints"),
        ], col_widths=(2.5, 4.0))

    # ── Section 3.7: Deployment ── own page so it doesn't stack on the table
    doc.add_page_break()
    section_heading(doc, "3.7  Deployment Architecture")
    body(doc,
         "SkillSphere is deployed on Vercel using a monorepo configuration. The React build "
         "output is served as static files from Vercel's global edge CDN. The Express "
         "application is wrapped as a Vercel Serverless Function invoked on each API request. "
         "Serverless functions can have a cold-start latency of 3–8 seconds on the free "
         "tier when the Node.js runtime has gone idle. SkillSphere handles this with two "
         "mechanisms: the MongoDB connection is cached via an isConnected flag so subsequent "
         "invocations skip reconnection overhead, and the Axios timeout is set to 12 seconds "
         "to absorb worst-case cold starts without returning a timeout error to the user.")

    buf7 = diag_deployment()
    insert_figure(doc, buf7, 5.8, "Figure 3.5 — Deployment Architecture Diagram")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────
    # CHAPTER 4: RESULTS AND DISCUSSION
    # ────────────────────────────────────────────────────────────
    heading(doc, "CHAPTER 4: RESULTS AND DISCUSSION", level=1)

    section_heading(doc, "4.1  Functional Testing Results")
    body(doc,
         "Thirteen functional test cases were executed manually against the deployed Vercel "
         "instance with a live MongoDB Atlas cluster and a Cloudinary account in free-tier mode. "
         "All thirteen passed. Table 4.1 lists each test case and its observed result.")
    spacer(doc, pts=6)
    add_2col_table(doc,
        ["Test Case", "Result"],
        [
            ("Register with an already-used email address", "400 Conflict — duplicate error returned correctly"),
            ("Login with incorrect password",               "401 Unauthorized — rejected without leaking stored hash"),
            ("JWT expires after 120 minutes",               "401 on next authenticated request; client redirected to login"),
            ("Course creation with AI outline",             "Form fields populated from Gemini response within 7 s"),
            ("Video upload to Cloudinary",                  "secure_url and public_id returned; temp file deleted"),
            ("Duplicate free-course enrollment",            "400 Already enrolled — second enrollment blocked"),
            ("PayPal payment order creation",               "approveUrl returned; Order document saved with 'pending' status"),
            ("Mark last lecture viewed",                    "completed flag set to true; completionDate recorded"),
            ("AI quiz generation (10 questions, 3 groups)", "Valid JSON groups returned; difficulty distribution respected"),
            ("Quiz attempt scoring — all correct",          "percentage = 100, passed = true"),
            ("AI chat streaming",                           "SSE chunks received in real time; 'done' event terminates stream"),
            ("Unenroll with progress cleanup",              "CourseProgress document deleted; StudentCourses updated"),
            ("Cloudinary duration backfill",                "Duration stored to Course document on first student access"),
        ], col_widths=(3.2, 3.3))

    section_heading(doc, "4.2  Performance Observations")
    body(doc,
         "Performance was measured by timing API calls under typical usage conditions — a "
         "single user session, free-tier Vercel and MongoDB Atlas, Gemini free quota. "
         "These numbers represent realistic conditions for a student project deployment, "
         "not a load-tested production environment.")
    add_2col_table(doc,
        ["Operation", "Observed Latency"],
        [
            ("Cold-start API response (first request after idle)",    "3 – 8 seconds"),
            ("Warm API response (subsequent requests)",               "100 – 300 ms"),
            ("AI course outline generation",                          "3 – 7 seconds"),
            ("AI quiz generation (20 questions across 2 groups)",     "5 – 12 seconds"),
            ("AI tutor chat — time to first token (SSE)",             "0.5 – 2 seconds"),
            ("Video upload to Cloudinary (50 MB file)",               "8 – 15 seconds"),
            ("MongoDB indexed query (e.g. findOne by email)",         "< 10 ms"),
            ("React SPA initial load (Vercel Edge CDN)",              "1 – 2 seconds"),
        ], col_widths=(3.5, 3.0))

    section_heading(doc, "4.3  Discussion")
    body(doc,
         "All thirteen test cases passed, confirming the core workflows are correct. What "
         "was more informative was what testing revealed about system behavior under "
         "realistic conditions.")
    body(doc,
         "AI response time was the most variable factor. Quiz generation ranged from 5 to "
         "12 seconds on identical inputs, depending on Gemini API load at the time of the "
         "request. There is nothing to tune at the application level to reduce this — cloud "
         "AI latency is external. What the application controls is the experience during "
         "that wait: the submit button disables on click and stays disabled, a loading "
         "indicator runs, and the state is clear to the user. No duplicate submissions "
         "occurred across all test runs.")
    body(doc,
         "Vercel cold starts were a real problem, not a theoretical one. The first API "
         "request after the function goes idle can take 3–8 seconds before the MongoDB "
         "connection is even established. The Axios timeout is set to 12 seconds specifically "
         "to absorb this. A standard 4-second timeout would fail on every cold start. Once "
         "warm, non-AI calls respond in under 300 ms and the latency goes unnoticed.")
    body(doc,
         "The most notable result came from the AI tutor's streaming design. First tokens "
         "arrive in 0.5–2 seconds, and the response builds on screen as text generates. "
         "During testing, even responses that took 9–10 seconds end-to-end did not feel slow "
         "because progress was visible from the first second. Replacing SSE with a "
         "spinner-and-wait design on the same backend would make those same responses feel "
         "roughly twice as long subjectively. SSE worked natively across every tested "
         "browser — Chrome, Firefox, Edge, Safari, and both mobile variants — with no "
         "polyfill required.")
    body(doc,
         "Table 4.3 compares SkillSphere's feature set against established platforms. No "
         "publicly available LMS combines AI outline generation, AI quiz generation, and a "
         "course-scoped streaming tutor in a single open-source, free-to-deploy package.")
    spacer(doc, pts=6)
    add_2col_table(doc,
        ["Feature", "SkillSphere   Moodle   Teachable   Coursera"],
        [
            ("AI Course Outline Generation",   "Yes          Plugin       No          No"),
            ("AI Quiz Generation",             "Yes          Plugin       No          No"),
            ("AI Tutor (Streaming)",           "Yes          No           No          No"),
            ("PayPal Payments",                "Yes          Plugin       Yes         No"),
            ("Free Deployment Tier",           "Yes (Vercel) Self-hosted  No ($39/mo) No"),
            ("Open Source",                    "Yes          Yes          No          No"),
        ], col_widths=(3.0, 3.5))

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────
    # CHAPTER 5: CONCLUSION
    # ────────────────────────────────────────────────────────────
    heading(doc, "CHAPTER 5: CONCLUSION", level=1)

    section_heading(doc, "5.1  Conclusion")
    body(doc,
         "The question this project started with was concrete: does building AI into an LMS "
         "architecture from the start produce a better outcome than adding it later as a "
         "plugin? After five months of building and testing SkillSphere: yes, measurably so.")
    body(doc,
         "The AI tutor is the clearest example. It works better than a general-purpose "
         "chatbot because it knows the course before the student opens the chat. The system "
         "instruction is built at request time from the actual course document — title, "
         "description, learning objectives — so the tutor answers in context rather than "
         "in generalities. No custom ML work was needed for this; just a well-structured "
         "prompt and a Gemini API call.")
    body(doc,
         "AI outline generation was the most surprising result in practice. The expectation "
         "was that it would be useful. What was unexpected was how an instructor went from a "
         "topic description to a reviewed and edited 12-lecture course structure in under "
         "20 minutes. Without the feature, that process takes most of a working day.")
    body(doc,
         "The MERN stack was the right choice for this type of application. JavaScript "
         "throughout eliminated the translation layer between database output and UI state. "
         "MongoDB documents fed directly into React form state with no serialization step. "
         "The document model suited course data well: a curriculum is inherently nested, and "
         "a single embedded-array document was simpler than any relational schema.")
    body(doc,
         "Vercel's serverless deployment worked, with documented caveats. Cold-start latency "
         "required explicit mitigation: the isConnected flag to reuse MongoDB connections "
         "across invocations, and the extended Axios timeout to prevent false error states "
         "on first load. The 10-second execution limit puts long AI requests close to the "
         "boundary on the hobby tier. For a student project with intermittent traffic, "
         "these trade-offs were acceptable. For a production deployment with sustained load, "
         "a persistent server process would be the better choice.")
    body(doc,
         "All thirteen test cases passed. The SSE streaming design made the AI tutor feel "
         "genuinely responsive despite real latency. The system does what it set out to do.")

    section_heading(doc, "5.2  Limitations")
    limitations = [
        "JWT tokens expire after 120 minutes with no refresh mechanism, requiring users to log in again for longer sessions.",
        "No email verification on registration — any email address, including non-existent ones, can be used to create an account.",
        "PayPal integration runs in sandbox mode only; switching to live payments requires a business account verification process beyond the scope of this project.",
        "Course browsing has no full-text search; filtering is limited to category, level, language, and price sort.",
        "No rate limiting on the authentication endpoints, leaving them theoretically vulnerable to automated credential stuffing.",
        "AI quiz quality can vary; the instructor review step is available but not enforced before publishing.",
    ]
    for lim in limitations:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(lim)
        set_font(run, size=12)
        para_spacing(p, before=2, after=3)

    section_heading(doc, "5.3  Future Scope")
    body(doc,
         "Several natural extensions would move SkillSphere from a functional prototype "
         "toward a production-grade platform:")
    future = [
        ("Short-term (3–6 months):",
         "JWT refresh token rotation, Nodemailer email verification, MongoDB Atlas Search "
         "for full-text course search, cursor-based pagination on the course list, and "
         "input validation using Joi or Zod across all API endpoints."),
        ("Medium-term (6–12 months):",
         "An instructor analytics dashboard showing completion rates and quiz performance "
         "aggregates, HLS adaptive bitrate streaming via Cloudinary's transcoding pipeline, "
         "course-level discussion forums, and a student rating and review system."),
        ("Long-term (12+ months):",
         "Personalised learning path recommendations driven by quiz performance history, "
         "automatic lecture transcript and summary generation using Gemini Vision on uploaded "
         "videos, a B2B multi-tenant mode for organisations to run private LMS instances, "
         "and an AI essay grader that scores free-text student submissions against a "
         "rubric provided by the instructor."),
    ]
    for term, desc in future:
        p = doc.add_paragraph()
        r1 = p.add_run(term + " ")
        set_font(r1, size=12, bold=True)
        r2 = p.add_run(desc)
        set_font(r2, size=12)
        para_spacing(p, before=4, after=4)

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────
    # REFERENCES
    # ────────────────────────────────────────────────────────────
    heading(doc, "REFERENCES")
    refs = [
        "[1] D. G. Oblinger, \"Education Nation: Six Leading Edges of Innovation in our Schools,\" EDUCAUSE Review, vol. 47, no. 3, 2012.",
        "[2] G. Siemens, \"Connectivism: A Learning Theory for the Digital Age,\" International Journal of Instructional Technology and Distance Learning, vol. 2, no. 1, pp. 3–10, 2005.",
        "[3] A. Vaswani et al., \"Attention Is All You Need,\" Advances in Neural Information Processing Systems, vol. 30, 2017.",
        "[4] T. Brown et al., \"Language Models are Few-Shot Learners,\" Advances in Neural Information Processing Systems, vol. 33, pp. 1877–1901, 2020.",
        "[5] M. Alario-Hoyos et al., \"Recommendations for the Design and Deployment of MOOCs,\" IEEE Transactions on Learning Technologies, vol. 10, no. 2, pp. 196–202, 2017.",
        "[6] S. Bhatt and J. Mudge, \"Machine Learning in Education: A Systematic Review,\" International Journal of Artificial Intelligence in Education, vol. 31, pp. 345–382, 2021.",
        "[7] A. Pinkwart, \"Another 25 Years of AIED? Challenges and Opportunities for Intelligent Educational Technologies of the Future,\" International Journal of Artificial Intelligence in Education, vol. 26, pp. 771–783, 2016.",
        "[8] M. V. Meza-Kubo and A. L. Morán, \"Technologies for Online Learning: A Systematic Review,\" Computers & Education, vol. 151, 2020.",
        "[9] MongoDB Inc., \"MongoDB Atlas Documentation,\" Available: https://www.mongodb.com/docs/atlas/",
        "[10] Google LLC, \"Gemini API Documentation — Generative AI,\" Available: https://ai.google.dev/gemini-api/docs",
        "[11] Cloudinary Inc., \"Cloudinary Node.js SDK Documentation,\" Available: https://cloudinary.com/documentation/node_integration",
        "[12] PayPal Inc., \"PayPal REST API Documentation,\" Available: https://developer.paypal.com/docs/api/overview/",
        "[13] Meta Open Source, \"React Documentation,\" Available: https://react.dev",
        "[14] Express.js Foundation, \"Express.js Documentation,\" Available: https://expressjs.com/",
        "[15] Mongoose ODM, \"Mongoose Documentation v7,\" Available: https://mongoosejs.com/docs/",
        "[16] Vercel Inc., \"Vercel Platform Documentation,\" Available: https://vercel.com/docs",
        "[17] M. Jones et al., \"JSON Web Token (JWT) — RFC 7519,\" Internet Engineering Task Force (IETF), May 2015.",
        "[18] W3C, \"Server-Sent Events Specification,\" Available: https://html.spec.whatwg.org/multipage/server-sent-events.html",
        "[19] Global Market Insights, \"E-Learning Market Size, Share & Trends Report 2023–2032,\" 2023.",
        "[20] R. Fernandez et al., \"Integration of Cloud Services for IoT-Based and Web Application Deployments,\" IEEE Internet of Things Journal, vol. 9, no. 5, pp. 3500–3512, 2022.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        set_font(run, size=11)
        para_spacing(p, before=2, after=5)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)

    doc.save(OUT)
    print(f"Thesis saved -> {OUT}")


if __name__ == "__main__":
    build_thesis()
